# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, color="00513B", sz="12", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_run(para, text, size_pt=10, bold=False, color_rgb=(30, 41, 59), font_name='맑은 고딕'):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    return run

doc = docx.Document()

# 1. A4 용지 및 1페이지 완벽 안착을 위한 여백 설정 (상하 12mm, 좌우 15mm)
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(0.47)    # 약 12mm
section.bottom_margin = Inches(0.47) # 약 12mm
section.left_margin = Inches(0.59)   # 약 15mm
section.right_margin = Inches(0.59)  # 약 15mm

# 2. 상단 발신/수신 바
p_header = doc.add_paragraph()
p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_header.paragraph_format.space_after = Pt(2)
add_run(p_header, "[Veterinary Exclusive · Wave 1 DM]  ", size_pt=9, bold=True, color_rgb=(0, 81, 59))
add_run(p_header, "발신: 에스앤제이(S&J) 동물병원 (대표: 홍순일 원장)\n수신: 전국 소동물병원 원장님 귀하", size_pt=9, bold=True, color_rgb=(71, 85, 105))

# 구분선
p_line = doc.add_paragraph()
p_line.paragraph_format.space_after = Pt(6)
add_run(p_line, "―" * 48, size_pt=10.5, bold=True, color_rgb=(0, 81, 59))

# 3. 문서 메인 타이틀
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(8)
p_title.paragraph_format.line_spacing = 1.15
add_run(p_title, "에스앤제이(S&J) 동물병원 학술마케팅부 공식 초청 서신\n", size_pt=14, bold=True, color_rgb=(15, 23, 42))
add_run(p_title, "“(전)서울대학교 수의과대학 출신 자문단과 약 6,000례 임상 경험을 기반으로 설계된\n수의전용 고미립 장내 유해물질 흡착·제거 보조 솔루션 MONSMECTA 공식 런칭”", size_pt=11.5, bold=True, color_rgb=(0, 81, 59))

# 4. 본문 편지글
p_body1 = doc.add_paragraph()
p_body1.paragraph_format.space_after = Pt(10)
p_body1.paragraph_format.line_spacing = 1.15
add_run(p_body1, "존경하는 원장님, 안녕하십니까.\n", size_pt=10, bold=True, color_rgb=(15, 23, 42))
add_run(p_body1, "진료 현장에서 파보·로타·코로나 등 중증 바이러스성 장염 및 설사 환축 치료 시, 장내 수분·독소·바이러스를 강력히 흡착해 설사를 멎게 하고 손상된 위장 점막을 효과적으로 보호하는 천연 점토(스멕타이트/몬모릴로나이트) 성분 흡착성 지사제 솔루션의 필요성을 잘 알고 계실 것입니다.\n이에 에스앤제이(S&J) 동물병원은 (전)서울대학교 수의과대학 출신 자문단과 함께, ", size_pt=10, color_rgb=(30, 41, 59))
add_run(p_body1, "약 6,000례의 임상 경험을 기반으로 설계된 소동물 전용 천연 고미립 장내 유해물질 흡착·제거 보조제 『몬스멕타(MONSMECTA)』", size_pt=10, bold=True, color_rgb=(0, 81, 59))
add_run(p_body1, "를 공식 런칭하고 전국 5,000개 동물병원 원장님들께 특별 우선 배포를 제안드립니다.", size_pt=10, color_rgb=(30, 41, 59))

# 5. 자문단 검증 요약 박스 (몬스멕타 Deep Emerald 테두리 & 소프트 에메랄드 배경)
tbl_adv = doc.add_table(rows=1, cols=1)
tbl_adv.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_adv = tbl_adv.cell(0, 0)
cell_adv.width = Inches(7.09)
set_cell_background(cell_adv, "F0FDF4") # 에메랄드 은은한 배경
set_cell_border(cell_adv, color="00513B", sz="12", val="single") # 선명한 몬스멕타 에메랄드 테두리 (1.5pt)
set_cell_margins(cell_adv, top=140, bottom=140, left=180, right=180)

p_adv_header = cell_adv.paragraphs[0]
p_adv_header.paragraph_format.space_after = Pt(6)
add_run(p_adv_header, "■ 바이러스학 분야 권위자 및 임상 자문단 검증 코멘트", size_pt=10, bold=True, color_rgb=(0, 81, 59))

adv_items = [
    ("[바이러스학] 박봉균 교수 (전 농림축산검역본부장, 바이러스학 분야 권위자): ", "“일반 지사제와 달리 1-Deoxynojirimycin(DNJ)의 여러 바이러스 억제 효과와 차별화된 장내 유해물질 흡착력은 임상 진료에서 매우 강력한 강점입니다.”"),
    ("[임상검증] 정성대 원장 (동진동물병원 원장): ", "“입자도, 수분 흡수율, 비표면적... 수의사가 진료실에서 처방할 때 가장 알고 싶어 하는 정확한 물성 데이터가 완벽하게 규명되었습니다.”"),
    ("[회복단축] 김동준 원장 (사랑동물병원 원장): ", "“정맥 수액(IV) 처치와 병행 투여 시, 장점막을 이중 보호하고 전해질 밸런스를 맞춰 파보 장염 환축의 증상 회복 기간 단축에 큰 도움을 줍니다.”")
]

for title, content in adv_items:
    p_item = cell_adv.add_paragraph()
    p_item.paragraph_format.space_after = Pt(4)
    p_item.paragraph_format.line_spacing = 1.12
    add_run(p_item, title, size_pt=9.5, bold=True, color_rgb=(15, 23, 42))
    add_run(p_item, content, size_pt=9.5, color_rgb=(51, 65, 85))

# 여백 문단 (박스 간 넉넉하고 편안한 공간 확보)
p_gap = doc.add_paragraph()
p_gap.paragraph_format.space_before = Pt(8)
p_gap.paragraph_format.space_after = Pt(8)

# 6. Wave 1 초도 물량 안내 및 QR코드 박스 (동일한 몬스멕타 Deep Emerald 톤 통일)
tbl_qr_box = doc.add_table(rows=1, cols=1)
tbl_qr_box.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_qr_box = tbl_qr_box.cell(0, 0)
cell_qr_box.width = Inches(7.09)
set_cell_background(cell_qr_box, "F0FDF4") # 몬스멕타 톤 일치
set_cell_border(cell_qr_box, color="00513B", sz="12", val="single") # 선명한 몬스멕타 에메랄드 테두리
set_cell_margins(cell_qr_box, top=140, bottom=140, left=160, right=160)

p_qr_title = cell_qr_box.paragraphs[0]
p_qr_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_qr_title.paragraph_format.space_after = Pt(6)
p_qr_title.paragraph_format.line_spacing = 1.15
add_run(p_qr_title, "[Wave 1 초도 양산 한정 선착순 우선 공급] ", size_pt=10, bold=True, color_rgb=(0, 81, 59))
add_run(p_qr_title, "전국 소동물병원 5,000처 중 단 170개 병원 선착순 배포\n", size_pt=10.5, bold=True, color_rgb=(15, 23, 42))
add_run(p_qr_title, "복잡한 주소 입력 없이 스마트폰 카메라로 아래 QR코드를 스캔하시면 1초 만에 접수 및 학술 데이터 열람이 가능합니다.", size_pt=9, color_rgb=(71, 85, 105))

# 내부 2열 QR코드 테이블
tbl_qrs = cell_qr_box.add_table(rows=1, cols=2)
tbl_qrs.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in tbl_qrs.rows:
    for c in row.cells:
        c.width = Inches(3.35)
        set_cell_background(c, "FFFFFF")
        set_cell_border(c, color="00513B", sz="6", val="single") # 서브 카드도 깔끔한 그린선
        set_cell_margins(c, top=100, bottom=100, left=100, right=100)

# 왼쪽 QR (발주)
c_left = tbl_qrs.cell(0, 0)
p_pic1 = c_left.paragraphs[0]
p_pic1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_pic1.paragraph_format.space_after = Pt(3)
if os.path.exists("monsmecta_qrcode_snj1.png"):
    p_pic1.add_run().add_picture("monsmecta_qrcode_snj1.png", width=Inches(0.95))
elif os.path.exists("public/assets/monsmecta_qrcode_snj1.png"):
    p_pic1.add_run().add_picture("public/assets/monsmecta_qrcode_snj1.png", width=Inches(0.95))

p_txt1 = c_left.add_paragraph()
p_txt1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_txt1.paragraph_format.space_after = Pt(0)
p_txt1.paragraph_format.line_spacing = 1.12
add_run(p_txt1, "[QR CODE 01]\n", size_pt=8.5, bold=True, color_rgb=(0, 81, 59))
add_run(p_txt1, "📦 초도 170세트 발주 및 특별 접수 QR\n", size_pt=9, bold=True, color_rgb=(15, 23, 42))
add_run(p_txt1, "카메라 스캔 시 즉시 병원 인증 및 접수", size_pt=8, color_rgb=(100, 116, 139))

# 오른쪽 QR (학술 데이터)
c_right = tbl_qrs.cell(0, 1)
p_pic2 = c_right.paragraphs[0]
p_pic2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_pic2.paragraph_format.space_after = Pt(3)
if os.path.exists("monsmecta_qrcode_snj2.png"):
    p_pic2.add_run().add_picture("monsmecta_qrcode_snj2.png", width=Inches(0.95))
elif os.path.exists("public/assets/monsmecta_qrcode_snj2.png"):
    p_pic2.add_run().add_picture("public/assets/monsmecta_qrcode_snj2.png", width=Inches(0.95))

p_txt2 = c_right.add_paragraph()
p_txt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_txt2.paragraph_format.space_after = Pt(0)
p_txt2.paragraph_format.line_spacing = 1.12
add_run(p_txt2, "[QR CODE 02]\n", size_pt=8.5, bold=True, color_rgb=(0, 81, 59))
add_run(p_txt2, "📊 자문단 학술 데이터 및 시나리오 QR\n", size_pt=9, bold=True, color_rgb=(15, 23, 42))
add_run(p_txt2, "(전)서울대 출신 자문단 데이터 및 상담 시나리오 열람", size_pt=8, color_rgb=(100, 116, 139))

# 7. 하단 푸터
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(8)
p_footer.paragraph_format.space_after = Pt(0)
p_footer.paragraph_format.line_spacing = 1.15
add_run(p_footer, "―" * 48 + "\n", size_pt=8.5, color_rgb=(203, 213, 225))
add_run(p_footer, "에스앤제이(S&J) 동물병원 | 대표: 홍순일 (닥터 젬스홍) | 사업자번호: 792-66-00615\n", size_pt=8.5, bold=True, color_rgb=(51, 65, 85))
add_run(p_footer, "주소: 경기도 용인시 처인구 포곡읍 선장1로 98-8 | 대표 연락처: 031-321-6562 | 공식 이메일: soonilhong@naver.com\n", size_pt=8.5, bold=True, color_rgb=(0, 81, 59))
add_run(p_footer, "※ 본 안내문은 전국 소동물병원 수의사 원장님께 한정하여 발송되는 수의학 전용 학술·처방 안내 문서입니다.", size_pt=7.5, color_rgb=(148, 163, 184))

out_docx_1 = os.path.abspath("public/assets/monsmecta_wave1-1_dm_letter_a4.docx")
out_docx_2 = os.path.abspath("public/assets/tournament_variant_16.docx")

doc.save(out_docx_1)
doc.save(out_docx_2)
print("Word 문서 수정 및 1등 당선작(Variant 16) 반영 완료!")

