import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), border_style.get('val', 'single'))
            node.set(qn('w:sz'), str(border_style.get('sz', 4)))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_run(p, text, size_pt=10, bold=False, italic=False, color_rgb=(30, 41, 59), font_name="맑은 고딕"):
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)
    return run

def create_orca_design(variant_id, primary_hex, secondary_hex, accent_hex, border_hex, font_title):
    doc = docx.Document()
    
    # A4 margins: 0.5 in top/bottom, 0.6 in left/right
    for s in doc.sections:
        s.top_margin = Inches(0.45)
        s.bottom_margin = Inches(0.45)
        s.left_margin = Inches(0.55)
        s.right_margin = Inches(0.55)
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_after = Pt(2)
    p_hdr.paragraph_format.line_spacing = 1.05
    
    tbl_top = doc.add_table(rows=1, cols=2)
    tbl_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = tbl_top.cell(0, 0), tbl_top.cell(0, 1)
    c1.width, c2.width = Inches(3.6), Inches(3.5)
    
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    add_run(p1, "Veterinary Exclusive · Wave 1 DM", size_pt=8.5, bold=True, color_rgb=(255, 255, 255), font_name=font_title)
    set_cell_background(c1, primary_hex)
    set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
    
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, "공문번호: S&J-2026-07-W1", size_pt=9, bold=True, color_rgb=(0, 81, 59), font_name=font_title)
    set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(4)
    add_run(p_title, "에스앤제이(S&J) 동물병원 학술마케팅부 공식 초청 서신", size_pt=15, bold=True, color_rgb=(15, 23, 42), font_name=font_title)
    
    tbl_meta = doc.add_table(rows=1, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    mc1, mc2 = tbl_meta.cell(0, 0), tbl_meta.cell(0, 1)
    mc1.width, mc2.width = Inches(3.55), Inches(3.55)
    set_cell_background(mc1, "F8FAFC")
    set_cell_background(mc2, "F8FAFC")
    set_cell_margins(mc1, top=80, bottom=80, left=100, right=100)
    set_cell_margins(mc2, top=80, bottom=80, left=100, right=100)
    
    mp1 = mc1.paragraphs[0]
    mp1.paragraph_format.space_after = Pt(0)
    add_run(mp1, "[발신] ", size_pt=8.5, bold=True, color_rgb=(0, 81, 59))
    add_run(mp1, "에스앤제이(S&J) 동물병원 (대표: 홍순일 원장)", size_pt=9, color_rgb=(30, 41, 59))
    
    mp2 = mc2.paragraphs[0]
    mp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mp2.paragraph_format.space_after = Pt(0)
    add_run(mp2, "[수신] ", size_pt=8.5, bold=True, color_rgb=(217, 119, 6))
    add_run(mp2, "전국 소동물병원 원장님 귀하", size_pt=9, bold=True, color_rgb=(0, 81, 59))
    
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(8)
    add_run(p_div, "―" * 52, size_pt=9, bold=True, color_rgb=(0, 81, 59))
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    p_sub.paragraph_format.line_spacing = 1.2
    add_run(p_sub, "“(전)서울대학교 수의과대학 출신 자문단과 약 6,000례 임상 경험을 기반으로 설계된\n수의전용 고미립 장내 유해물질 흡착·제거 보조 솔루션 MONSMECTA 공식 런칭”", size_pt=11.5, bold=True, color_rgb=(0, 81, 59), font_name=font_title)
    
    p_body = doc.add_paragraph()
    p_body.paragraph_format.space_after = Pt(6)
    p_body.paragraph_format.line_spacing = 1.18
    add_run(p_body, "존경하는 원장님, 안녕하십니까.\n", size_pt=9.5, bold=True, color_rgb=(15, 23, 42))
    add_run(p_body, "진료 현장에서 파보·로타·코로나 등 중증 바이러스성 장염 및 설사 환축 치료 시, 장내 수분·독소·바이러스를 강력히 흡착해 설사를 멎게 하고 손상된 위장 점막을 효과적으로 보호하는 천연 점토(스멕타이트/몬모릴로나이트) 성분 흡착성 지사제 솔루션의 필요성을 잘 알고 계실 것입니다.\n", size_pt=9.5, color_rgb=(30, 41, 59))
    add_run(p_body, "이에 에스앤제이(S&J) 동물병원은 (전)서울대학교 수의과대학 출신 자문단과 함께, ", size_pt=9.5, color_rgb=(30, 41, 59))
    add_run(p_body, "약 6,000례의 임상 경험을 기반으로 설계된 소동물 전용 천연 고미립 장내 유해물질 흡착·제거 보조제 『몬스멕타(MONSMECTA)』", size_pt=9.5, bold=True, color_rgb=(0, 81, 59))
    add_run(p_body, "를 공식 런칭하고 전국 5,000개 동물병원 원장님들께 특별 우선 배포를 제안드립니다.", size_pt=9.5, color_rgb=(30, 41, 59))
    
    # 자문단 박스
    tbl_adv = doc.add_table(rows=1, cols=1)
    tbl_adv.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_adv = tbl_adv.cell(0, 0)
    c_adv.width = Inches(7.1)
    set_cell_background(c_adv, "F8FAFC")
    set_cell_margins(c_adv, top=100, bottom=100, left=140, right=140)
    set_cell_borders(c_adv, left={'val': 'single', 'sz': 12, 'color': primary_hex})
    
    p_adv_h = c_adv.paragraphs[0]
    p_adv_h.paragraph_format.space_after = Pt(4)
    add_run(p_adv_h, "■ 바이러스학 분야 권위자 및 임상 자문단 검증 코멘트", size_pt=9.5, bold=True, color_rgb=(0, 81, 59), font_name=font_title)
    
    adv_items = [
        ("[바이러스학] 박봉균 교수 (전 농림축산검역본부장, 바이러스학 분야 권위자): ", "“일반 지사제와 달리 1-Deoxynojirimycin(DNJ)의 여러 바이러스 억제 효과와 차별화된 장내 유해물질 흡착력은 임상 진료에서 매우 강력한 강점입니다.”"),
        ("[임상검증] 정성대 원장 (동진동물병원 원장): ", "“입자도, 수분 흡수율, 비표면적... 수의사가 진료실에서 처방할 때 가장 알고 싶어 하는 정확한 물성 데이터가 완벽하게 규명되었습니다.”"),
        ("[회복단축] 김동준 원장 (사랑동물병원 원장): ", "“정맥 수액(IV) 처치와 병행 투여 시, 장점막을 이중 보호하고 전해질 밸런스를 맞춰 파보 장염 환축의 증상 회복 기간 단축에 큰 도움을 줍니다.”")
    ]
    for title, content in adv_items:
        pa = c_adv.add_paragraph()
        pa.paragraph_format.space_after = Pt(3)
        pa.paragraph_format.line_spacing = 1.12
        add_run(pa, title, size_pt=8.5, bold=True, color_rgb=(15, 23, 42))
        add_run(pa, content, size_pt=8.5, color_rgb=(71, 85, 105))
        
    # QR 박스
    tbl_qr = doc.add_table(rows=1, cols=1)
    tbl_qr.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_qr = tbl_qr.cell(0, 0)
    c_qr.width = Inches(7.1)
    set_cell_background(c_qr, "FFFBEB")
    set_cell_margins(c_qr, top=90, bottom=90, left=120, right=120)
    set_cell_borders(c_qr, top={'val': 'single', 'sz': 6, 'color': accent_hex},
                           bottom={'val': 'single', 'sz': 6, 'color': accent_hex},
                           left={'val': 'single', 'sz': 6, 'color': accent_hex},
                           right={'val': 'single', 'sz': 6, 'color': accent_hex})
                           
    pq_t = c_qr.paragraphs[0]
    pq_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pq_t.paragraph_format.space_after = Pt(3)
    add_run(pq_t, "[Wave 1 초도 양산 한정 선착순 우선 공급] ", size_pt=9, bold=True, color_rgb=(217, 119, 6))
    add_run(pq_t, "전국 소동물병원 5,000처 중 단 170개 병원 선착순 배포\n", size_pt=10, bold=True, color_rgb=(15, 23, 42))
    add_run(pq_t, "스마트폰 카메라로 아래 QR코드를 스캔하시면 1초 만에 접수 및 학술 데이터 열람이 가능합니다.", size_pt=8.5, color_rgb=(71, 85, 105))
    
    # 2열 QR
    t_qrs = c_qr.add_table(rows=1, cols=2)
    t_qrs.alignment = WD_TABLE_ALIGNMENT.CENTER
    cq1, cq2 = t_qrs.cell(0, 0), t_qrs.cell(0, 1)
    cq1.width, cq2.width = Inches(3.35), Inches(3.35)
    set_cell_background(cq1, "FFFFFF")
    set_cell_background(cq2, "FFFFFF")
    set_cell_margins(cq1, top=60, bottom=60, left=60, right=60)
    set_cell_margins(cq2, top=60, bottom=60, left=60, right=60)
    
    p_img1 = cq1.paragraphs[0]
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.paragraph_format.space_after = Pt(2)
    if os.path.exists("public/assets/monsmecta_qrcode_snj1.png"):
        p_img1.add_run().add_picture("public/assets/monsmecta_qrcode_snj1.png", width=Inches(0.85))
    p_t1 = cq1.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_after = Pt(0)
    add_run(p_t1, "[QR 01] 초도 170세트 발주 접수\n", size_pt=8.5, bold=True, color_rgb=(0, 81, 59))
    add_run(p_t1, "스캔 시 즉시 병원 인증 및 접수", size_pt=7.5, color_rgb=(100, 116, 139))
    
    p_img2 = cq2.paragraphs[0]
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.paragraph_format.space_after = Pt(2)
    if os.path.exists("public/assets/monsmecta_qrcode_snj2.png"):
        p_img2.add_run().add_picture("public/assets/monsmecta_qrcode_snj2.png", width=Inches(0.85))
    p_t2 = cq2.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t2.paragraph_format.space_after = Pt(0)
    add_run(p_t2, "[QR 02] 자문단 학술 데이터 및 시나리오\n", size_pt=8.5, bold=True, color_rgb=(217, 119, 6))
    add_run(p_t2, "(전)서울대 수의대 자문 데이터 열람", size_pt=7.5, color_rgb=(100, 116, 139))
    
    # 푸터
    p_ftr = doc.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ftr.paragraph_format.space_before = Pt(8)
    p_ftr.paragraph_format.space_after = Pt(0)
    p_ftr.paragraph_format.line_spacing = 1.15
    add_run(p_ftr, "에스앤제이(S&J) 동물병원\n", size_pt=9.5, bold=True, color_rgb=(15, 23, 42), font_name=font_title)
    add_run(p_ftr, "대표: 홍순일 (닥터 젬스홍) | 사업자번호: 792-66-00615 | 주소: 경기도 용인시 처인구 포곡읍 선장1로 98-8\n", size_pt=8, color_rgb=(71, 85, 105))
    add_run(p_ftr, "대표 연락처: 031-321-6562 | 공식 이메일: soonilhong@naver.com\n", size_pt=8.5, bold=True, color_rgb=(0, 81, 59))
    add_run(p_ftr, "※ 본 안내문은 전국 소동물병원 수의사 원장님께 한정하여 발송되는 수의학 전용 학술·처방 안내 문서입니다.", size_pt=7.5, color_rgb=(148, 163, 184))
    
    filename_docx = f"public/assets/tournament_variant_{variant_id}.docx"
    doc.save(filename_docx)
    return filename_docx

# 16개 에이전트 토너먼트 테마 정의
agents_theme = [
    (1, "00513B", "003D2B", "F59E0B", "FCD34D", "맑은 고딕"),
    (2, "0F172A", "1E293B", "D97706", "FDE68A", "맑은 고딕"),
    (3, "1E3A8A", "1E40AF", "0284C7", "BAE6FD", "맑은 고딕"),
    (4, "065F46", "047857", "059669", "A7F3D0", "맑은 고딕"),
    (5, "831843", "9D174D", "BE185D", "FBCFE8", "맑은 고딕"),
    (6, "4C1D95", "5B21B6", "7C3AED", "DDD6FE", "맑은 고딕"),
    (7, "7C2D12", "9A3412", "C2410C", "FFEDD5", "맑은 고딕"),
    (8, "111827", "1F2937", "374151", "E5E7EB", "맑은 고딕"),
    (9, "00513B", "064E3B", "10B981", "D1FAE5", "맑은 고딕"),
    (10, "1E1B4B", "312E81", "4338CA", "C7D2FE", "맑은 고딕"),
    (11, "312E81", "3730A3", "4F46E5", "E0E7FF", "맑은 고딕"),
    (12, "064E3B", "065F46", "047857", "A7F3D0", "맑은 고딕"),
    (13, "701A75", "86198F", "A21CAF", "F5D0FE", "맑은 고딕"),
    (14, "881337", "9F1239", "BE123C", "FECDD3", "맑은 고딕"),
    (15, "134E4A", "115E59", "0D9488", "99F6E4", "맑은 고딕"),
    (16, "0F172A", "00513B", "D97706", "FCD34D", "맑은 고딕")
]

for var_id, p_hex, s_hex, a_hex, b_hex, f_name in agents_theme:
    create_orca_design(var_id, p_hex, s_hex, a_hex, b_hex, f_name)
    print(f"Generated variant {var_id}")
